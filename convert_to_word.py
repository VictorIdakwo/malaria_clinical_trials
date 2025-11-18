"""
Convert SYSTEM_OVERVIEW.md to Microsoft Word (.docx)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_heading(doc, text, level=1):
    """Add a formatted heading to the document"""
    # Remove markdown symbols
    text = text.replace('#', '').strip()
    text = text.replace('**', '')
    
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph to the document"""
    if not text.strip():
        return
    
    paragraph = doc.add_paragraph()
    
    # Handle bold text
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italic text
            italic_parts = re.split(r'(\*.*?\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                    run = paragraph.add_run(ipart[1:-1])
                    run.italic = True
                else:
                    # Handle inline code
                    code_parts = re.split(r'(`.*?`)', ipart)
                    for cpart in code_parts:
                        if cpart.startswith('`') and cpart.endswith('`'):
                            run = paragraph.add_run(cpart[1:-1])
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                        else:
                            paragraph.add_run(cpart)
    
    return paragraph

def add_code_block(doc, code_text):
    """Add a code block to the document"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    paragraph.style = 'No Spacing'
    
    # Add shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    paragraph._element.get_or_add_pPr().append(shading_elm)
    
    return paragraph

def add_table_from_markdown(doc, lines):
    """Convert markdown table to Word table"""
    # Parse table
    rows = []
    for line in lines:
        if '|' in line and not line.strip().startswith('|--'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_data.replace('**', '')
            
            # Bold header row
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

def convert_markdown_to_word(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue
        
        # Handle tables
        if '|' in line and not line.strip().startswith('#'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            add_table_from_markdown(doc, table_lines)
            in_table = False
            table_lines = []
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            add_heading(doc, line, level=min(level, 3))
        
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)
        
        # Handle bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            paragraph = add_paragraph(doc, text)
            paragraph.style = 'List Bullet'
        
        # Handle numbered lists
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            paragraph = add_paragraph(doc, text)
            paragraph.style = 'List Number'
        
        # Handle regular paragraphs
        elif line.strip():
            add_paragraph(doc, line.strip())
        
        # Add blank line for spacing
        else:
            doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"[SUCCESS] Successfully converted to: {docx_file}")

if __name__ == "__main__":
    md_file = "SYSTEM_OVERVIEW.md"
    docx_file = "SYSTEM_OVERVIEW.docx"
    
    print(f"Converting {md_file} to {docx_file}...")
    convert_markdown_to_word(md_file, docx_file)
    print("Done! [SUCCESS]")
